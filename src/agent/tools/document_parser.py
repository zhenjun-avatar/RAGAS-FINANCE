"""Minimal document parser for the node-centric ingestion pipeline."""

from __future__ import annotations

import asyncio
import re
from html.parser import HTMLParser
from pathlib import Path
from typing import Any

import chardet
import fitz
import pdfplumber
from docx import Document

SUPPORTED_FORMATS = {"pdf", "txt", "docx", "html", "htm"}


async def parse_document_content(file_path: str, file_type: str) -> dict[str, Any]:
    loop = asyncio.get_running_loop()
    result = await loop.run_in_executor(None, _parse_document_content_sync, file_path, file_type)
    if result.get("success"):
        text = result.get("text") or ""
        result["word_count"] = len(text)
    return result


def _parse_document_content_sync(file_path: str, file_type: str) -> dict[str, Any]:
    path = Path(file_path)
    normalized_type = _normalize_file_type(file_type or path.suffix)

    if not path.exists():
        return {"success": False, "error": f"文件不存在：{file_path}"}
    if normalized_type not in SUPPORTED_FORMATS:
        supported = ", ".join(sorted(SUPPORTED_FORMATS))
        return {"success": False, "error": f"仅支持以下文件类型：{supported}"}

    try:
        if normalized_type == "pdf":
            text, metadata = _parse_pdf(path)
        elif normalized_type == "txt":
            text, metadata = _parse_txt(path)
        elif normalized_type in ("html", "htm"):
            text, metadata = _parse_html(path)
        else:
            text, metadata = _parse_docx(path)
    except Exception as exc:
        return {"success": False, "error": f"文档解析失败：{exc}"}

    cleaned = _clean_text(text)
    if not cleaned:
        return {"success": False, "error": "解析结果为空"}
    metadata.update(
        {
            "file_type": normalized_type,
            "file_name": path.name,
            "file_size": path.stat().st_size,
            "char_count": len(cleaned),
        }
    )
    return {"success": True, "text": cleaned, "metadata": metadata}


def _normalize_file_type(file_type: str) -> str:
    return (file_type or "").lower().replace(".", "").strip()


def _clean_text(text: str) -> str:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    lines = [line.strip() for line in normalized.split("\n")]
    return "\n".join(line for line in lines if line).strip()


def _parse_pdf(path: Path) -> tuple[str, dict[str, Any]]:
    page_texts: list[str] = []
    page_count = 0

    with fitz.open(path) as pdf:
        page_count = len(pdf)
        for index, page in enumerate(pdf, start=1):
            text = _clean_text(page.get_text("text"))
            if text:
                page_texts.append(f"## 第 {index} 页\n\n{text}")

    if not page_texts:
        with pdfplumber.open(path) as pdf:
            page_count = len(pdf.pages)
            for index, page in enumerate(pdf.pages, start=1):
                text = _clean_text(page.extract_text() or "")
                if text:
                    page_texts.append(f"## 第 {index} 页\n\n{text}")

    return "\n\n".join(page_texts), {"page_count": page_count, "parser": "pdf"}


def _parse_txt(path: Path) -> tuple[str, dict[str, Any]]:
    raw = path.read_bytes()
    detection = chardet.detect(raw)
    encoding = detection.get("encoding") or "utf-8"
    text = raw.decode(encoding, errors="ignore")
    return text, {"encoding": encoding, "parser": "txt"}


class _VisibleTextHTMLParser(HTMLParser):
    """Convert HTML into structured plain text for chunking.

    Keeps section boundaries (headings/paragraphs/lists) and linearizes table rows.
    """

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self._lines: list[str] = []
        self._current_text: list[str] = []
        self._suppress = 0
        self._in_table = 0
        self._row_cells: list[str] = []
        self._cell_text: list[str] = []
        self._in_cell = 0
        self._heading_level: int | None = None

    def _normalize_ws(self, text: str) -> str:
        return re.sub(r"\s+", " ", text).strip()

    def _flush_current_text(self, *, blank_after: bool = False) -> None:
        text = self._normalize_ws("".join(self._current_text))
        self._current_text = []
        if text:
            self._lines.append(text)
            if blank_after:
                self._lines.append("")

    def _flush_table_row(self) -> None:
        if self._in_cell:
            return
        cells = [self._normalize_ws(x) for x in self._row_cells if self._normalize_ws(x)]
        self._row_cells = []
        if cells:
            self._lines.append(f"TABLE | {' | '.join(cells)}")

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in ("script", "style", "noscript"):
            self._suppress += 1
            return
        if self._suppress > 0:
            return
        low = tag.lower()
        if low in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            self._flush_current_text(blank_after=True)
            self._heading_level = int(low[1])
        elif low in {"p", "div", "section", "article", "ul", "ol", "li"}:
            self._flush_current_text(blank_after=True)
        elif low == "br":
            self._flush_current_text(blank_after=False)
        elif low == "table":
            self._flush_current_text(blank_after=True)
            self._in_table += 1
        elif low == "tr":
            self._flush_current_text(blank_after=False)
            self._flush_table_row()
            self._row_cells = []
        elif low in {"td", "th"}:
            self._in_cell += 1
            self._cell_text = []

    def handle_endtag(self, tag: str) -> None:
        if tag in ("script", "style", "noscript") and self._suppress > 0:
            self._suppress -= 1
            return
        if self._suppress > 0:
            return
        low = tag.lower()
        if low in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            heading_text = self._normalize_ws("".join(self._current_text))
            self._current_text = []
            if heading_text:
                level = self._heading_level or 2
                self._lines.append(f"{'#' * min(6, max(1, level))} {heading_text}")
                self._lines.append("")
            self._heading_level = None
        elif low in {"p", "div", "section", "article", "ul", "ol", "li"}:
            self._flush_current_text(blank_after=True)
        elif low in {"td", "th"} and self._in_cell > 0:
            cell = self._normalize_ws("".join(self._cell_text))
            self._cell_text = []
            self._in_cell -= 1
            if cell:
                self._row_cells.append(cell)
        elif low == "tr":
            self._flush_table_row()
            self._lines.append("")
        elif low == "table" and self._in_table > 0:
            self._flush_table_row()
            self._in_table -= 1
            self._lines.append("")

    def handle_data(self, data: str) -> None:
        if self._suppress > 0 or not data:
            return
        if self._in_cell > 0:
            self._cell_text.append(data)
        else:
            self._current_text.append(data)

    def to_text(self) -> str:
        self._flush_current_text(blank_after=False)
        self._flush_table_row()
        out: list[str] = []
        blank_pending = False
        for raw in self._lines:
            line = raw.strip()
            if not line:
                if out:
                    blank_pending = True
                continue
            if blank_pending:
                out.append("")
                blank_pending = False
            out.append(line)
        return "\n".join(out).strip()


def _parse_html(path: Path) -> tuple[str, dict[str, Any]]:
    raw = path.read_bytes()
    markup: str | None = None
    used_enc = "utf-8"
    for enc in ("utf-8", "latin-1"):
        try:
            markup = raw.decode(enc)
            used_enc = enc
            break
        except UnicodeDecodeError:
            continue
    if markup is None:
        markup = raw.decode("utf-8", errors="replace")
    parser = _VisibleTextHTMLParser()
    try:
        parser.feed(markup)
        parser.close()
    except Exception:
        pass
    text = parser.to_text()
    table_row_count = text.count("\nTABLE | ") + (1 if text.startswith("TABLE | ") else 0)
    return text, {"parser": "html", "encoding": used_enc, "table_row_count": table_row_count}


def _parse_docx(path: Path) -> tuple[str, dict[str, Any]]:
    doc = Document(path)
    blocks: list[str] = []

    for paragraph in doc.paragraphs:
        text = _clean_text(paragraph.text)
        if text:
            blocks.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [_clean_text(cell.text) for cell in row.cells]
            row_text = " | ".join(cell for cell in cells if cell)
            if row_text:
                blocks.append(row_text)

    return "\n\n".join(blocks), {"parser": "docx"}
